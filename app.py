import streamlit as st
from datetime import datetime, timedelta
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# =========================================================
# 基礎設定
# =========================================================
PROVIDER_NAME = "高如慧"
BANK_NAME = "中國信託商業銀行"
BANK_CODE = "822"
ACCOUNT_NUMBER = "783540208870"
PHASE2_TUTORIAL_URL = "https://youtu.be/caoZAO8tyNs"

# =========================================================
# Page config
# =========================================================
st.set_page_config(
    page_title="廣告投放工具",
    page_icon="📝",
    layout="centered"
)

st.title("📝 廣告投放合作工具")
st.caption("合約生成 × 啟動前資料蒐集")
st.markdown("---")

# =========================================================
# Session state（僅存結果，不做流程 gating）
# =========================================================
for k, v in {
    "client_message": "",
    "payment_message": "",
    "docx_bytes": b"",
    "last_party_a_name": "",
    "phase2_message": "",
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# =========================================================
# Word 字型設定
# =========================================================
def set_run_font(run, size=12, bold=False):
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

# =========================================================
# Word 合約生成
# =========================================================
def generate_docx_bytes(party_a, payment_opt, start_dt, pay_day, pay_dt):
    doc = Document()
    style = doc.styles["Normal"]
    style.paragraph_format.line_spacing = 1.5

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("廣告投放服務合約書")
    set_run_font(r, 18, True)
    doc.add_paragraph("")

    if payment_opt == "17,000元/月（每月付款）":
        end_dt = start_dt + timedelta(days=30)
        period_text = f"自 {start_dt:%Y 年 %m 月 %d 日} 起至 {end_dt:%Y 年 %m 月 %d 日} 止，共 1 個月。"
        price_text = "新台幣壹萬柒仟元整（NT$17,000）／月"
    else:
        end_dt = start_dt + timedelta(days=90)
        period_text = f"自 {start_dt:%Y 年 %m 月 %d 日} 起至 {end_dt:%Y 年 %m 月 %d 日} 止，共 3 個月。"
        price_text = "新台幣肆萬伍仟元整（NT$45,000）／三個月"

    p = doc.add_paragraph()
    set_run_font(p.add_run(f"甲方：{party_a}\n乙方：{PROVIDER_NAME}\n\n"))
    set_run_font(p.add_run(f"合約期間：{period_text}\n"))
    set_run_font(p.add_run(f"服務費用：{price_text}\n"))

    doc.add_paragraph("\n（以下略，依你原本合約條款）")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# =========================================================
# Sidebar 導覽（只切畫面）
# =========================================================
with st.sidebar:
    st.header("導覽")
    nav = st.radio(
        "選擇階段",
        ["第一階段｜合約", "第二階段｜啟動前確認"]
    )

# =========================================================
# 第一階段｜合約
# =========================================================
if nav == "第一階段｜合約":

    st.header("📄 合約生成")

    payment_option = st.radio(
        "付款方案",
        ["17,000元/月（每月付款）", "45,000元/三個月（一次付款）"]
    )

    start_date = st.date_input(
        "合作啟動日",
        value=datetime.now().date() + timedelta(days=7)
    )

    payment_day = st.slider("每月付款日", 1, 28, 5) if "月" in payment_option else None
    payment_date = start_date - timedelta(days=3) if "三個月" in payment_option else None

    party_a_name = st.text_input("甲方名稱")

    if st.button("📝 生成 Word 合約", type="primary"):
        if not party_a_name.strip():
            st.error("請填寫甲方名稱")
        else:
            st.session_state.docx_bytes = generate_docx_bytes(
                party_a_name, payment_option, start_date, payment_day, payment_date
            )
            st.session_state.last_party_a_name = party_a_name

            st.session_state.client_message = f"""【合約確認】
甲方：{party_a_name}
乙方：{PROVIDER_NAME}
方案：{payment_option}
啟動日：{start_date}
"""

            st.session_state.payment_message = f"""【收款資訊】
銀行：{BANK_NAME} ({BANK_CODE})
帳號：{ACCOUNT_NUMBER}
"""

            st.success("合約已生成")

    if st.session_state.docx_bytes:
        st.subheader("📤 請客戶回傳")
        st.code(st.session_state.client_message)
        st.code(st.session_state.payment_message)

        st.download_button(
            "⬇️ 下載 Word 合約",
            st.session_state.docx_bytes,
            file_name=f"合約_{st.session_state.last_party_a_name}.docx"
        )

# =========================================================
# 第二階段｜啟動前確認（永遠可進）
# =========================================================
else:
    st.header("🚀 第二階段｜啟動前確認 & 資料蒐集")
    st.caption("📌 資料可分次填寫，不需要一次完成")

    st.video(PHASE2_TUTORIAL_URL)

    st.subheader("✅ 確認事項（照實勾選即可）")
    c1, c2 = st.columns(2)
    with c1:
        ad_account = st.checkbox("廣告帳號已開啟")
        pixel = st.checkbox("像素事件已埋放")
    with c2:
        fanpage = st.checkbox("粉專已建立")
        bm = st.checkbox("企業管理平台已建立")

    st.subheader("🧾 須提供事項")
    fanpage_url = st.text_input("粉專網址")
    landing_url = st.text_input("廣告導向頁")

    st.markdown("**競爭對手粉專（最多三個）**")
    comp1 = st.text_input("競品 1")
    comp2 = st.text_input("競品 2")
    comp3 = st.text_input("競品 3")

    who_problem = st.text_area("你的產品/服務要解決誰的問題？")
    what_problem = st.text_area("要解決什麼問題？")
    how_solve = st.text_area("如何解決？")
    budget = st.text_input("第一個月預算")

    def status(v): return "✅ 已完成" if v else "⬜ 未完成"

    if st.button("📌 產生回傳內容", type="primary"):
        msg = f"""【第二階段啟動資料】
甲方：{st.session_state.last_party_a_name or "（未填）"}

【確認事項】
- 廣告帳號：{status(ad_account)}
- 像素事件：{status(pixel)}
- 粉專：{status(fanpage)}
- BM：{status(bm)}

【資料】
- 粉專網址：{fanpage_url or "（未填）"}
- 導向頁：{landing_url or "（未填）"}

【競品】
1) {comp1 or "—"}
2) {comp2 or "—"}
3) {comp3 or "—"}

【定位】
- 對象：{who_problem or "—"}
- 問題：{what_problem or "—"}
- 解法：{how_solve or "—"}

【首月預算】
- {budget or "—"}
"""
        st.session_state.phase2_message = msg
        st.success("回傳內容已生成")

    if st.session_state.phase2_message:
        st.subheader("📤 請客戶複製回傳")
        st.code(st.session_state.phase2_message)
